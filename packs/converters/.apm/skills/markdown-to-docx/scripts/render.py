#!/usr/bin/env python3
"""
render.py — fill a branded Word template from a Markdown artifact.

This is the deterministic renderer for the `markdown-to-docx` skill. The agent
assembles the content model and invokes this script; the script fills a
user-provided `.docx` template at its Jinja tags (via docxtpl) rather than
building a document from scratch, so the designer's cover page, styles, headers,
and placed logo survive.

Verbs
    --check                 import-probe docxtpl; exit 0 (present) / 2 (absent)
    inspect <template>      print the template's declared Jinja variables
    render  <markdown>      fill a template and write a .docx
        --template <path>   pre-tagged .docx to fill; omit for the opt-out default
        --output <path>     output path (default: <markdown-basename>.docx in CWD)

Stdout markers (for agent parsing)
    FILLPOINTS: <variable>  (inspect — one per declared template variable)
    GUIDANCE: <msg>         template carries no Jinja tags; how to add them
    OUTPUT: <path>          path to the written .docx
    FILLED: <n>             count of template variables a value was supplied for
    WARNING: <msg>          non-fatal issue, surface to the user

Trust model
    A user-supplied template is trusted-author input, consistent with the
    converters pack's local-files-trusted stance. Because a `.docx` template
    carries Jinja2 *source*, a malicious template author could embed
    server-side template injection (SSTI); this is an accepted, out-of-scope
    risk for a trusted-author template, as are XXE / zip-bomb on a crafted
    archive. Two cheap controls hold regardless of template trust and are
    enforced here: docxtpl `autoescape=True` (user content is XML-escaped, not
    interpolated) and output-path confinement (the output path is assembled
    from model-influenced content).

Errors go to stderr; exit code 1 on failure, 2 on a missing library.
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path
from typing import NamedTuple

PIP_INSTALL = "python -m pip install 'docxtpl>=0.16.0'"

GUIDANCE_NO_TAGS = (
    "this .docx carries no Jinja fill-points, so there is nothing to fill. Add "
    "tags in Word where you want content: `{{ title }}` for a scalar, a "
    "`{%p for it in items %}{{ it }}{%p endfor %}` paragraph loop for a list, and "
    "a `{%tr for r in rows %}` row loop for a table. Author each tag inside one "
    "uniform run (don't let Word split `{{`/`}}` across runs by autocorrect or "
    "mid-tag formatting). Then re-run. The skill never silently converts an "
    "untagged document — that would discard your brand."
)


class RenderResult(NamedTuple):
    written: bool
    filled: int
    warnings: list[str]
    guidance: str | None


# --------------------------------------------------------------------------- #
# Output-path confinement
# --------------------------------------------------------------------------- #
def confine(path: Path, root: Path) -> Path:
    """Resolve ``path`` (following symlinks) and require it under ``root``.

    Path-*component* containment (``root`` is the resolved path or among its
    ``.parents``), not a string-prefix check, so a sibling like ``workdir-evil``
    is rejected against root ``workdir``. Raises ``ValueError`` on escape.
    """
    resolved = path.resolve()
    root_resolved = root.resolve()
    if resolved == root_resolved or root_resolved in resolved.parents:
        return resolved
    raise ValueError(
        f"path {path!r} resolves to {resolved} which is outside the working "
        f"directory {root_resolved}; refusing to read/write outside the project"
    )


# --------------------------------------------------------------------------- #
# Markdown → Jinja context (the map step)
# --------------------------------------------------------------------------- #
def build_context(text: str) -> dict:
    """Project a Markdown artifact onto a docxtpl Jinja context.

    Mapping: front-matter ``key: value`` lines → top-level scalars
    (for ``{{ key }}``); the first list → ``items`` (for a ``{%p for %}``
    paragraph loop); the first Markdown table → ``rows`` as a list of
    ``{column: value}`` dicts (for a ``{%tr for %}`` row loop). A ``sections``
    list mirroring the document structure is always provided.

    Pure transform — no I/O.
    """
    scalars: dict[str, str] = {}
    sections: list[dict] = []
    lines = text.splitlines()
    i = 0

    if lines and lines[0].strip() == "---":
        j = 1
        while j < len(lines) and lines[j].strip() != "---":
            raw = lines[j]
            if ":" in raw:
                key, _, value = raw.partition(":")
                scalars[key.strip()] = value.strip().strip("\"'")
            j += 1
        i = j + 1 if j < len(lines) else len(lines)

    current: dict | None = None

    def cell_split(row: str) -> list[str]:
        return [c.strip() for c in row.strip().strip("|").split("|")]

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            current = {"heading": stripped[level:].strip(), "bullets": [], "table": None}
            sections.append(current)
            i += 1
            continue

        if current is not None and (stripped.startswith("- ") or stripped.startswith("* ")):
            current["bullets"].append(stripped[2:].strip())
            i += 1
            continue

        if (
            current is not None
            and stripped.startswith("|")
            and i + 1 < len(lines)
            and "-" in lines[i + 1]
            and set(lines[i + 1].strip().replace("|", "").replace(":", "").strip()) <= {"-", " "}
            and lines[i + 1].strip().replace("|", "").strip()
        ):
            header = cell_split(stripped)
            body_rows = []
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = cell_split(lines[i])
                body_rows.append({header[k]: (cells[k] if k < len(cells) else "") for k in range(len(header))})
                i += 1
            current["table"] = {"header": header, "rows": body_rows}
            continue

        i += 1

    # Top-level convenience handles for the common single-list / single-table doc.
    items = next((s["bullets"] for s in sections if s["bullets"]), [])
    rows = next((s["table"]["rows"] for s in sections if s["table"]), [])

    context = dict(scalars)
    context["sections"] = sections
    context["items"] = items
    context["rows"] = rows
    return context


# --------------------------------------------------------------------------- #
# docxtpl helpers (imported lazily so --check runs without the library)
# --------------------------------------------------------------------------- #
def inspect_template(template: Path) -> list[str]:
    """Return the sorted set of undeclared Jinja variables in the template."""
    from docxtpl import DocxTemplate

    doc = DocxTemplate(str(template))
    return sorted(doc.get_undeclared_template_variables())


def render_docx(text: str, template: Path, output: Path) -> RenderResult:
    """Fill ``template`` from the Markdown in ``text``; write ``output``.

    If the template carries no Jinja tags, returns guidance instead of writing
    (never a silent convert). User content is escaped via ``autoescape=True``.
    """
    from docxtpl import DocxTemplate

    doc = DocxTemplate(str(template))
    variables = doc.get_undeclared_template_variables()
    if not variables:
        return RenderResult(written=False, filled=0, warnings=[], guidance=GUIDANCE_NO_TAGS)

    context = build_context(text)
    # autoescape is OFF by default in docxtpl; turn it ON so user content is
    # XML-escaped rather than injected as raw markup into the document.
    doc.render(context, autoescape=True)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))

    # FILLED counts variables a *non-empty* value was supplied for — the
    # always-present convenience handles (items/rows/sections) don't inflate it.
    filled = sum(1 for v in variables if context.get(v))
    warnings = []
    unfilled = variables - set(context.keys())
    if unfilled:
        warnings.append(
            "template declares variables with no Markdown source: "
            + ", ".join(sorted(unfilled))
        )
    return RenderResult(written=True, filled=filled, warnings=warnings, guidance=None)


def render_default_docx(text: str, output: Path) -> RenderResult:
    """Opt-out path: write the Markdown into a bare python-docx document.

    Used only when the user explicitly declines a branded template. The result
    is unbranded — front-matter scalars become paragraphs, headings become
    Word headings, lists become bullets, the first table becomes a Word table.
    Not the primary path: branding requires a template (``render_docx``).
    """
    from docx import Document

    context = build_context(text)
    doc = Document()
    title = context.get("title")
    if title:
        doc.add_heading(str(title), level=0)
    for key, value in context.items():
        if key in ("title", "items", "rows", "sections"):
            continue
        doc.add_paragraph(f"{key}: {value}")

    for section in context.get("sections", []):
        if section.get("heading"):
            doc.add_heading(section["heading"], level=1)
        for bullet in section.get("bullets", []):
            doc.add_paragraph(bullet, style="List Bullet")
        table = section.get("table")
        if table:
            header, body = table["header"], table["rows"]
            t = doc.add_table(rows=1, cols=len(header))
            for c, h in enumerate(header):
                t.rows[0].cells[c].text = h
            for row in body:
                cells = t.add_row().cells
                for c, h in enumerate(header):
                    cells[c].text = str(row.get(h, ""))

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    return RenderResult(
        written=True,
        filled=0,
        warnings=["no template supplied — produced an UNBRANDED .docx via python-docx"],
        guidance=None,
    )


# --------------------------------------------------------------------------- #
# CLI
# --------------------------------------------------------------------------- #
def cmd_check() -> int:
    try:
        import docxtpl  # noqa: F401
    except ImportError:
        print(f"docxtpl is not installed. Run:\n    {PIP_INSTALL}", file=sys.stderr)
        return 2
    return 0


def cmd_inspect(args) -> int:
    root = Path.cwd()
    template = confine(Path(args.template), root)
    if not template.is_file():
        print(f"ERROR: template not found: {args.template}", file=sys.stderr)
        return 1
    variables = inspect_template(template)
    if not variables:
        print(f"GUIDANCE: {GUIDANCE_NO_TAGS}")
        return 0
    for v in variables:
        print(f"FILLPOINTS: {v}")
    return 0


def cmd_render(args) -> int:
    root = Path.cwd()
    md_path = confine(Path(args.markdown), root)
    if not md_path.is_file():
        print(f"ERROR: markdown not found: {args.markdown}", file=sys.stderr)
        return 1

    output = Path(args.output) if args.output else Path(md_path.stem + ".docx")
    output = confine(output, root)
    text = md_path.read_text(encoding="utf-8")

    if not args.template:
        # Explicit opt-out (the agent omits --template only after the user
        # declines a branded template, per SKILL.md): bare python-docx document.
        result = render_default_docx(text, output)
    else:
        template = confine(Path(args.template), root)
        if not template.is_file():
            print(f"ERROR: template not found: {args.template}", file=sys.stderr)
            return 1
        result = render_docx(text, template, output)
    if result.guidance:
        print(f"GUIDANCE: {result.guidance}")
        return 0
    for w in result.warnings:
        print(f"WARNING: {w}")
    print(f"OUTPUT: {output}")
    print(f"FILLED: {result.filled}")
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Fill a Word template from Markdown.")
    parser.add_argument("--check", action="store_true", help="probe docxtpl; exit 0/2")
    sub = parser.add_subparsers(dest="verb")

    p_inspect = sub.add_parser("inspect", help="print the template's Jinja variables")
    p_inspect.add_argument("template")

    p_render = sub.add_parser("render", help="fill a template and write a .docx")
    p_render.add_argument("markdown")
    p_render.add_argument("--template", help="pre-tagged .docx to fill")
    p_render.add_argument("--output", help="output path (default: <basename>.docx in CWD)")
    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)

    if args.check:
        return cmd_check()

    rc = cmd_check()
    if rc != 0:
        return rc

    try:
        if args.verb == "inspect":
            return cmd_inspect(args)
        if args.verb == "render":
            return cmd_render(args)
    except ValueError as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1

    parser.print_help()
    return 1


if __name__ == "__main__":
    raise SystemExit(main())
