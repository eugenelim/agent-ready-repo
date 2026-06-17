"""TDD suite for the markdown-to-docx renderer.

Covers the contracts the spec names: Jinja-variable enumeration, Markdown→
context mapping, the no-fill-points guidance, the Tier-1 `--check` probe,
`autoescape=True` on user content, and output-path confinement — plus an
end-to-end render that re-opens the produced .docx and asserts the filled
values are present.

The tagged fixture template is built with python-docx at test time rather than
committed as an opaque binary; the builder below is its authorship. Run with
`python -m pytest` from this directory.
"""

from __future__ import annotations

import subprocess
import sys
import zipfile
from pathlib import Path

import pytest

import render

HERE = Path(__file__).resolve().parent

MARKDOWN = """\
---
title: Q3 Review
author: Finance
---

# Highlights

- Revenue up 12%
- Two new markets

# Numbers

| Metric | Value |
| --- | --- |
| ARR | 4.2M |
| Churn | 1.1% |
"""


def _build_tagged_template(path: Path) -> Path:
    """Author a Jinja-tagged .docx template (the fill-point authorship)."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("{{ title }}")
    doc.add_paragraph("{%p for it in items %}")
    doc.add_paragraph("{{ it }}")
    doc.add_paragraph("{%p endfor %}")
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "{%tr for r in rows %}"
    table.cell(1, 0).text = "{{ r.Metric }}"
    table.cell(1, 1).text = "{{ r.Value }}"
    table.cell(2, 0).text = "{%tr endfor %}"
    doc.save(str(path))
    return path


def _build_untagged_template(path: Path) -> Path:
    from docx import Document

    doc = Document()
    doc.add_paragraph("A plain document with no Jinja tags.")
    doc.save(str(path))
    return path


# --------------------------------------------------------------------------- #
# inspect — declared Jinja variables
# --------------------------------------------------------------------------- #
def test_inspect_returns_undeclared_variables(tmp_path):
    template = _build_tagged_template(tmp_path / "template.docx")
    assert render.inspect_template(template) == ["items", "rows", "title"]


# --------------------------------------------------------------------------- #
# map — Markdown → Jinja context
# --------------------------------------------------------------------------- #
def test_build_context_projects_scalars_list_and_table_rows():
    context = render.build_context(MARKDOWN)

    # front-matter → {{ var }} scalars
    assert context["title"] == "Q3 Review"
    assert context["author"] == "Finance"

    # a list → a {%p for %} loop context
    assert context["items"] == ["Revenue up 12%", "Two new markets"]

    # a Markdown table → a {%tr for %} row context (list of {column: value} dicts)
    assert context["rows"] == [
        {"Metric": "ARR", "Value": "4.2M"},
        {"Metric": "Churn", "Value": "1.1%"},
    ]


# --------------------------------------------------------------------------- #
# no-fill-points guidance
# --------------------------------------------------------------------------- #
def test_empty_markdown_parses_without_crash():
    assert render.build_context("") == {"items": [], "rows": [], "sections": []}


def test_unfilled_template_variable_is_warned(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("{{ title }}")
    doc.add_paragraph("{{ author_signature }}")  # declared, no Markdown source
    doc.save(str(tmp_path / "t.docx"))

    out = tmp_path / "out.docx"
    result = render.render_docx("---\ntitle: Q3\n---\n", tmp_path / "t.docx", out)
    assert result.written is True
    assert any("author_signature" in w for w in result.warnings)


def test_filled_count_ignores_empty_convenience_handles(tmp_path):
    # A template declaring `rows` but a Markdown doc with no table must not
    # report rows as filled (the convenience handle defaults to []).
    from docx import Document

    doc = Document()
    doc.add_paragraph("{{ title }}")
    table = doc.add_table(rows=3, cols=1)
    table.cell(0, 0).text = "{%tr for r in rows %}"
    table.cell(1, 0).text = "{{ r.x }}"
    table.cell(2, 0).text = "{%tr endfor %}"
    doc.save(str(tmp_path / "t.docx"))

    out = tmp_path / "out.docx"
    result = render.render_docx("---\ntitle: Q3\n---\n", tmp_path / "t.docx", out)
    assert result.filled == 1  # title only; empty `rows` is not counted


def test_render_default_writes_unbranded_docx(tmp_path):
    from docx import Document

    out = tmp_path / "out.docx"
    result = render.render_default_docx(MARKDOWN, out)
    assert result.written is True
    assert any("UNBRANDED" in w for w in result.warnings)
    assert out.is_file()
    blob = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "Q3 Review" in blob and "Revenue up 12%" in blob


def test_untagged_template_yields_guidance_not_silent_convert(tmp_path):
    template = _build_untagged_template(tmp_path / "plain.docx")
    assert render.inspect_template(template) == []

    out = tmp_path / "out.docx"
    result = render.render_docx(MARKDOWN, template, out)
    assert result.written is False
    assert result.guidance and "fill-point" in result.guidance
    assert not out.exists(), "an untagged template must not be silently converted"


# --------------------------------------------------------------------------- #
# --check — Tier-1 probe
# --------------------------------------------------------------------------- #
def test_check_exit_zero_when_present():
    assert render.cmd_check() == 0


def test_check_exit_two_when_absent(monkeypatch):
    monkeypatch.setitem(sys.modules, "docxtpl", None)
    assert render.cmd_check() == 2


# --------------------------------------------------------------------------- #
# autoescape — user content is escaped, not interpolated
# --------------------------------------------------------------------------- #
def test_user_content_is_xml_escaped(tmp_path):
    template = _build_tagged_template(tmp_path / "template.docx")
    md = "---\ntitle: <b>R&D</b>\n---\n\n# Heading\n\n- item\n"
    out = tmp_path / "out.docx"
    result = render.render_docx(md, template, out)
    assert result.written is True

    document_xml = zipfile.ZipFile(str(out)).read("word/document.xml").decode("utf-8")
    # autoescape=True turns the `<`/`&` into entities; an off-by-default render
    # would inject `<b>` as raw markup. Assert the escaped form is present and
    # the metacharacters were not interpolated as live XML.
    assert "&lt;b&gt;R&amp;D&lt;/b&gt;" in document_xml


# --------------------------------------------------------------------------- #
# confinement
# --------------------------------------------------------------------------- #
def test_confine_rejects_parent_traversal(tmp_path):
    root = tmp_path / "workdir"
    root.mkdir()
    with pytest.raises(ValueError):
        render.confine(root / ".." / "evil.docx", root)


def test_confine_rejects_symlink_escape(tmp_path):
    root = tmp_path / "workdir"
    root.mkdir()
    outside = tmp_path / "outside"
    outside.mkdir()
    link = root / "link"
    link.symlink_to(outside, target_is_directory=True)
    with pytest.raises(ValueError):
        render.confine(link / "evil.docx", root)


def test_confine_rejects_sibling_prefix(tmp_path):
    root = tmp_path / "workdir"
    root.mkdir()
    (tmp_path / "workdir-evil").mkdir()
    with pytest.raises(ValueError):
        render.confine(tmp_path / "workdir-evil" / "out.docx", root)


# --------------------------------------------------------------------------- #
# end-to-end render (manual-QA mode, exercised as an integration test)
# --------------------------------------------------------------------------- #
def test_render_fills_template_and_reopened_values_present(tmp_path):
    from docx import Document

    _build_tagged_template(tmp_path / "template.docx")
    (tmp_path / "report.md").write_text(MARKDOWN, encoding="utf-8")

    result = subprocess.run(
        [sys.executable, str(HERE / "render.py"), "render", "report.md",
         "--template", "template.docx", "--output", "out.docx"],
        cwd=tmp_path, capture_output=True, text=True,
    )
    assert result.returncode == 0, result.stderr
    assert "OUTPUT:" in result.stdout
    out = tmp_path / "out.docx"
    assert out.is_file()

    doc = Document(str(out))
    paras = "\n".join(p.text for p in doc.paragraphs)
    assert "Q3 Review" in paras
    assert "Revenue up 12%" in paras

    cells = [c.text for t in doc.tables for r in t.rows for c in r.cells]
    assert "ARR" in cells and "4.2M" in cells
